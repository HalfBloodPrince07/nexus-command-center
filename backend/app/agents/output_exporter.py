"""Phase 6 — OutputExporter: export synthesized report to PDF and DOCX."""
from __future__ import annotations

import logging
import shutil
import subprocess
from pathlib import Path
from typing import AsyncGenerator

from backend.config import settings

logger = logging.getLogger(__name__)


def _md_path(slug: str) -> Path:
    return settings.DEEP_RESEARCH_DIR / slug / "report.md"


def _output_dir(slug: str) -> Path:
    return settings.DEEP_RESEARCH_DIR / slug


# ── PDF export ────────────────────────────────────────────────────────────────

def export_to_pdf(md_path: Path, output_dir: Path, title: str = "Research Report") -> Path | None:
    pdf_path = output_dir / "report.pdf"

    # Strategy 1: pandoc + xelatex
    if shutil.which("pandoc") and shutil.which("xelatex"):
        try:
            subprocess.run(
                [
                    "pandoc", str(md_path),
                    "-o", str(pdf_path),
                    "--pdf-engine=xelatex",
                    "--toc",
                    "--toc-depth=3",
                    "--highlight-style=tango",
                    "-V", "geometry:margin=1in",
                    "-V", "fontsize=11pt",
                    "-V", "linestretch=1.4",
                    "-V", f"title={title}",
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            logger.info("PDF exported via pandoc+xelatex: %s", pdf_path)
            return pdf_path
        except subprocess.CalledProcessError as exc:
            logger.warning("pandoc+xelatex failed: %s", exc.stderr[:300])

    # Strategy 2: pandoc with wkhtmltopdf
    if shutil.which("pandoc") and shutil.which("wkhtmltopdf"):
        try:
            subprocess.run(
                [
                    "pandoc", str(md_path),
                    "-o", str(pdf_path),
                    "--pdf-engine=wkhtmltopdf",
                    "--toc",
                ],
                check=True,
                capture_output=True,
            )
            logger.info("PDF exported via pandoc+wkhtmltopdf: %s", pdf_path)
            return pdf_path
        except subprocess.CalledProcessError as exc:
            logger.warning("pandoc+wkhtmltopdf failed: %s", exc.stderr[:300] if exc.stderr else "")

    # Strategy 3: weasyprint (pure Python)
    try:
        import markdown as md_lib
        from weasyprint import HTML

        source = md_path.read_text(encoding="utf-8")
        html_body = md_lib.markdown(source, extensions=["tables", "fenced_code", "toc"])
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: Georgia, serif; max-width: 800px; margin: auto;
         font-size: 11pt; line-height: 1.6; padding: 2em; }}
  h1 {{ color: #1a1a2e; border-bottom: 2px solid #16213e; padding-bottom: 0.3em; }}
  h2 {{ color: #16213e; margin-top: 2em; }}
  h3 {{ color: #0f3460; }}
  blockquote {{ border-left: 4px solid #e74c3c; background: #fdf2f2; padding: 0.5em 1em; }}
  code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-size: 0.9em; }}
  pre code {{ display: block; padding: 1em; overflow-x: auto; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  td, th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
  th {{ background: #1a1a2e; color: white; }}
  tr:nth-child(even) {{ background: #f9f9f9; }}
</style>
</head><body>{html_body}</body></html>"""
        HTML(string=full_html).write_pdf(str(pdf_path))
        logger.info("PDF exported via weasyprint: %s", pdf_path)
        return pdf_path
    except ImportError:
        logger.debug("weasyprint not available for PDF export")
    except Exception as exc:
        logger.warning("weasyprint PDF export failed: %s", exc)

    logger.warning("No PDF exporter available for slug %s", md_path.parent.name)
    return None


# ── DOCX export ───────────────────────────────────────────────────────────────

def _docx_template() -> str | None:
    template = Path(__file__).parent.parent.parent / "templates" / "research_template.docx"
    return str(template) if template.exists() else None


def export_to_docx(md_path: Path, output_dir: Path) -> Path | None:
    docx_path = output_dir / "report.docx"

    # Strategy 1: pandoc
    if shutil.which("pandoc"):
        cmd = [
            "pandoc", str(md_path),
            "-o", str(docx_path),
            "--toc",
            "--highlight-style=tango",
        ]
        tmpl = _docx_template()
        if tmpl:
            cmd += ["--reference-doc", tmpl]
        try:
            subprocess.run(cmd, check=True, capture_output=True)
            logger.info("DOCX exported via pandoc: %s", docx_path)
            return docx_path
        except subprocess.CalledProcessError as exc:
            logger.warning("pandoc DOCX failed: %s", exc.stderr[:300] if exc.stderr else "")

    # Strategy 2: python-docx fallback
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        source = md_path.read_text(encoding="utf-8")
        doc = Document()
        _md_to_docx(doc, source)
        doc.save(str(docx_path))
        logger.info("DOCX exported via python-docx: %s", docx_path)
        return docx_path
    except ImportError:
        logger.debug("python-docx not available for DOCX export")
    except Exception as exc:
        logger.warning("python-docx DOCX export failed: %s", exc)

    logger.warning("No DOCX exporter available for slug %s", md_path.parent.name)
    return None


def _md_to_docx(doc: object, source: str) -> None:
    """Minimal markdown → python-docx conversion for the fallback path."""
    from docx.shared import Pt, RGBColor

    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("| "):
            # Skip table rows (too complex for minimal impl)
            pass
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:], style="Quote")
        elif stripped == "---":
            doc.add_page_break()
        elif stripped:
            doc.add_paragraph(stripped)


# ── Agent class ───────────────────────────────────────────────────────────────

class OutputExporterAgent:
    async def run(
        self,
        slug: str,
        report_md: str | None = None,
        title: str = "Research Report",
    ) -> AsyncGenerator[dict, None]:
        """
        Export the report at DEEP_RESEARCH_DIR/{slug}/report.md to PDF and DOCX.
        If report_md is provided, write it to disk first (re-saves the report).
        """
        out_dir = _output_dir(slug)
        md_file = _md_path(slug)

        if report_md:
            out_dir.mkdir(parents=True, exist_ok=True)
            md_file.write_text(report_md, encoding="utf-8")

        if not md_file.exists():
            yield {
                "type": "error",
                "agent": "Exporter",
                "detail": f"report.md not found at {md_file} — cannot export",
            }
            return

        output_paths: dict[str, str] = {"md": str(md_file)}

        # PDF
        yield {"type": "progress", "agent": "Exporter", "stage": "pdf", "detail": "Exporting to PDF..."}
        pdf = export_to_pdf(md_file, out_dir, title=title)
        if pdf:
            output_paths["pdf"] = str(pdf)

        # DOCX
        yield {"type": "progress", "agent": "Exporter", "stage": "docx", "detail": "Exporting to DOCX..."}
        docx = export_to_docx(md_file, out_dir)
        if docx:
            output_paths["docx"] = str(docx)

        yield {
            "type": "result",
            "agent": "Exporter",
            "slug": slug,
            "output_paths": output_paths,
            "formats_exported": list(output_paths.keys()),
        }


output_exporter = OutputExporterAgent()
